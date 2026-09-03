from __future__ import annotations

import json
import urllib.request
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "documentation_assets" / "screens"
OUTPUT = ROOT / "Документация_системы_расписания.docx"
API = "http://127.0.0.1:8080/api"

NAVY = "14213D"
BLUE = "4F46E5"
HEADING_BLUE = "2E74B5"
HEADING_DARK = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "0E7A53"
AMBER = "8A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "1A1A1A"

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def fetch(path: str):
    with urllib.request.urlopen(API + path, timeout=20) as response:
        return json.load(response)


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=None, bold=None, italic=None, color=BLACK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    assert sum(widths_dxa) == TABLE_WIDTH_DXA
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D0D5DD", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, HEADING_BLUE, 18, 10),
        ("Heading 2", 13, HEADING_BLUE, 14, 7),
        ("Heading 3", 12, HEADING_DARK, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)


def add_custom_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_bullet = max(existing_abstract or [0]) + 1
    abstract_decimal = abstract_bullet + 1
    num_bullet = max(existing_num or [0]) + 1
    num_decimal = num_bullet + 1

    def abstract(abstract_id, fmt, text):
        node = OxmlElement("w:abstractNum")
        node.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        node.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        node.append(lvl)
        numbering.append(node)

    def num(num_id, abstract_id):
        node = OxmlElement("w:num")
        node.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        node.append(abstract_ref)
        numbering.append(node)

    abstract(abstract_bullet, "bullet", "•")
    abstract(abstract_decimal, "decimal", "%1.")
    num(num_bullet, abstract_bullet)
    num(num_decimal, abstract_decimal)
    return num_bullet, num_decimal


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def add_list_item(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_run_font(lead, bold=True)
        tail = p.add_run(text[len(bold_prefix):])
        set_run_font(tail)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_label_paragraph(doc, label, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    lead = p.add_run(label + " ")
    set_run_font(lead, bold=True, color=HEADING_DARK)
    body = p.add_run(text)
    set_run_font(body)
    return p


def add_callout(doc, title, text, fill=CALLOUT, accent=BLUE):
    def decorate(paragraph, before, after):
        paragraph.paragraph_format.left_indent = Inches(0.12)
        paragraph.paragraph_format.right_indent = Inches(0.08)
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        p_pr.append(shd)
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "20")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), accent)
        borders.append(left)
        p_pr.append(borders)

    p = doc.add_paragraph()
    decorate(p, 6, 1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=accent)
    p2 = doc.add_paragraph()
    decorate(p2, 0, 8)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5)


def add_table(doc, headers, rows, widths_dxa, font_size=9.2, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(header))
        set_run_font(run, size=font_size, bold=True, color=HEADING_DARK)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2 == 1:
                set_cell_shading(cells[index], "FAFBFC")
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if index > 0 and len(str(value)) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_screenshot(doc, filename, caption, width=6.3):
    path = ASSETS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_with_next = False


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("СИСТЕМА РАСПИСАНИЯ  |  ПРЕЗЕНТАЦИОННАЯ ДОКУМЕНТАЦИЯ")
    set_run_font(hr, size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def bool_text(value):
    return "Включено" if value else "Выключено"


def value_text(value):
    if isinstance(value, bool):
        return bool_text(value)
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


ROOT_ADVICE = {
    "solver_time_limit_seconds": "Резерв для монолитного режима; обычно не трогать.",
    "week_time_limit_seconds": "30–120 с для работы; увеличивать при сложных неделях.",
    "solver_workers": "Оставить 4 на текущем ПК.",
    "solver_max_memory_mb": "Ниже доступной RAM; не завышать.",
    "stop_after_first_solution": "Включать только для быстрого черновика.",
    "linearization_level": "Обычно 0; 1–2 тестировать только по замерам.",
    "symmetry_level": "2 — нормальный рабочий вариант.",
    "random_seed": "Менять, если поиск застрял на одной конфигурации.",
    "quality_improvement_seconds": "0 — черновик, 5–10 — работа, 30 — финал.",
    "use_quality_objective": "Включать для финальной вычитки и уменьшения штрафов.",
    "hard_no_student_windows": "Осторожно: может сделать неделю неразрешимой.",
    "hard_no_teacher_windows": "Использовать только при достаточной свободе сетки.",
    "hard_min_study_days_per_week": "Сначала держать мягким ограничением.",
    "hard_min_2_teacher_pairs_per_day": "Очень жёстко; обычно выключено.",
    "strict_all_theory_before_labs": "Включать только при обязательном методическом требовании.",
    "optimize_teacher_windows": "Для финального прогона; требует quality objective.",
    "optimize_student_windows": "Основной мягкий механизм сокращения окон.",
    "group_week_missing_day_weight": "Больше значение — сильнее стремление заполнить учебные дни.",
    "subject_missing_bucket_weight": "0 отключает приоритет равномерности по периодам.",
    "subject_bucket_overload_weight": "Повышать, если дисциплина собирается комком.",
    "subject_missing_segment_weight": "Повышать, если предмет исчезает в начале/конце семестра.",
    "student_five_pair_day_weight": "Повышать, чтобы реже ставить по 5 пар.",
    "student_late_slot_weight": "0 быстрее; повышение сдвигает занятия раньше.",
    "teacher_late_slot_weight": "Обычно 0; включать по отдельному требованию.",
    "teacher_window_weight": "Работает только при оптимизации окон преподавателей.",
    "student_window_weight": "Рабочий диапазон 200–600.",
    "min_student_pairs_per_study_day": "1–2; значение 2 уменьшает одиночные приезды.",
    "max_student_pairs_per_day": "Жёсткий предел; обычно 5.",
    "min_student_study_days_per_week": "Цель, а при HARD — обязательное число дней.",
    "subject_spread_bucket_available_days": "12 примерно соответствует двум учебным неделям.",
    "min_subject_spread_total_slots": "9999 фактически отключает spread; около 4 включает.",
    "subject_bucket_extra_slots": "Больше значение делает равномерность мягче.",
    "subject_bucket_min_capacity": "Обычно 2; менять вместе со spread.",
    "normal_subject_active_bucket_unit": "Техническая настройка spread; обычно 2.",
    "block_subject_active_bucket_unit": "Техническая настройка УП; обычно 4.",
    "min_initial_theory_slots_before_labs": "Минимум теории до первой лабораторной.",
}


def build():
    data = fetch("/data")
    audit = fetch("/audit")
    hours = fetch("/hours")
    solver = fetch("/settings/solver-config")
    try:
        schedule = fetch("/schedule")
    except Exception:
        schedule = {"groups": []}

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    bullet_num, decimal_num = add_custom_numbering(doc)
    doc.core_properties.title = "Документация системы автоматического формирования расписания"
    doc.core_properties.subject = "Краткое презентационное руководство по функционалу и ROOT-настройкам"
    doc.core_properties.author = "Проект системы расписания"
    doc.core_properties.keywords = "расписание, OR-Tools, учёт часов, техникум"

    # Cover: editorial_cover pattern + compact_reference_guide tokens.
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("ПРЕЗЕНТАЦИОННОЕ РУКОВОДСТВО")
    set_run_font(kr, size=10, bold=True, color=BLUE)
    kicker.paragraph_format.space_after = Pt(18)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Система автоматического\nформирования расписания")
    set_run_font(tr, size=29, bold=True, color=NAVY)
    title.paragraph_format.space_after = Pt(10)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Функционал, выполненные работы и параметры решателя ROOT")
    set_run_font(sr, size=14, color=HEADING_DARK)
    subtitle.paragraph_format.space_after = Pt(40)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Версия для демонстрации  •  28 августа 2026")
    set_run_font(mr, size=11, bold=True, color=MUTED)
    meta.paragraph_format.space_after = Pt(36)
    add_callout(
        doc,
        "Цель продукта",
        "Собрать данные учебного плана, автоматически сформировать допустимое расписание, контролировать часы и публиковать понятную версию для студентов.",
        fill="EEF2FF",
        accent=BLUE,
    )
    doc.add_page_break()

    doc.add_heading("1. Кратко о текущем результате", level=1)
    add_label_paragraph(
        doc,
        "Что это:",
        "единая локальная веб-система для диспетчера техникума: справочники, генерация OR-Tools, ручная корректировка, контроль часов, замены и публикация расписания.",
    )
    add_label_paragraph(
        doc,
        "Главный принцип:",
        "исходная нагрузка приходит из «Вклеек», расписание расставляет занятия по датам и парам, а учёт часов показывает план, поставлено, фактический зачёт и остаток.",
    )
    add_callout(
        doc,
        "Статус",
        "Фронтенд и серверная часть собираются локально. Генератор, импорт, отчёты и конфликтный сценарий аудиторий покрыты автоматическим smoke-тестом.",
        fill="ECFDF3",
        accent=GREEN,
    )

    stats = [
        ("Группы", len(data.get("groups", []))),
        ("Преподаватели", len(data.get("teachers", []))),
        ("Занятия / нагрузки", len(data.get("lessons", []))),
        ("Аудитории", len(data.get("rooms", []))),
        ("Типы аудиторий", len(data.get("room_types", []))),
        ("Группы в сгенерированном расписании", len(schedule.get("groups", []))),
    ]
    add_table(doc, ["Объект", "Текущее количество"], stats, [6800, 2560], font_size=9.5)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    nr = note.add_run("Снимок данных на момент подготовки документации.")
    set_run_font(nr, size=8.5, italic=True, color=MUTED)

    doc.add_heading("2. Какие работы проведены", level=1)
    work_rows = [
        ("Сборка и запуск", "Собраны frontend и C++ backend, настроен локальный запуск и демонстрационный доступ."),
        ("Тестовые данные", "Подключены ранее сгенерированные данные и рабочий просмотр расписания."),
        ("Импорт Excel", "Добавлен импорт «Вклеек» с выбором семестра, предпросмотром, предупреждениями и аудитом до применения."),
        ("Надёжность данных", "Стабильные UID, резервные версии, история изменений и откат до одной из последних 50 версий."),
        ("Рабочее время", "Периоды, дни и диапазоны пар для групп и преподавателей; массовое назначение выбранным объектам."),
        ("Преподаватели", "Закреплённый кабинет, приоритет площадки, описание ответственности за кабинет, индивидуальная недоступность."),
        ("Учёт часов", "План / поставлено / зачтено / остаток, недели, даты пар, состав нагрузки, фильтры и контроль перегруза."),
        ("Замены", "Перенос фактически зачтённых часов с заболевшего преподавателя на заменяющего по конкретной дате и паре."),
        ("Аудитории", "Фонд кабинетов, вместимость, оборудование, корпус и редактируемые типы 1/2/3."),
        ("Автозамена кабинета", "При конфликте подбирается свободная совместимая аудитория; причина и замена попадают в отчёт."),
        ("Экспорт и публикация", "Excel, PDF, студенческая публикация, универсальный учёт часов и Excel по структуре старого образца."),
        ("Решатель", "Недельная генерация, ограничение ресурсов, быстрый поиск первого решения, фаза улучшения качества и диагностические отчёты."),
        ("Проверка", "Автотест фиктивной сетки проверяет аудит, генерацию, часы, замены, публикацию и конфликт кабинетов."),
    ]
    add_table(doc, ["Направление", "Результат"], work_rows, [2500, 6860], font_size=9.0)

    doc.add_heading("3. Рекомендуемый рабочий цикл", level=1)
    workflow = [
        "Задать даты семестра и проверить основные настройки.",
        "Импортировать «Вклейки» через предпросмотр, не заменяя данные вслепую.",
        "Проверить аудит: ошибки исправить, предупреждения осознанно принять.",
        "Уточнить группы, преподавателей, рабочее время, площадки и аудиторный фонд.",
        "Запустить генерацию с нуля либо с фиксацией ручного/предыдущего варианта.",
        "Проверить расписание, отчёт аудиторий, окна и выполнение часов.",
        "При необходимости скорректировать сетку в Конструкторе и сохранить.",
        "Опубликовать студенческую версию и выгрузить Excel/PDF/учёт часов.",
    ]
    for item in workflow:
        add_list_item(doc, item, decimal_num)
    add_callout(
        doc,
        "Важно",
        "После изменения входных данных или ROOT-параметров необходимо запускать регенерацию. Публикация не происходит автоматически.",
        fill="FFF8E8",
        accent=AMBER,
    )
    doc.add_page_break()

    doc.add_heading("4. Разделы приложения", level=1)
    doc.add_heading("4.1. Стартовый экран и режим для студентов", level=2)
    add_label_paragraph(doc, "Стартовый экран:", "разделяет вход на режим «Студентам» и рабочую область диспетчера.")
    add_label_paragraph(doc, "Студенческий режим:", "показывает только опубликованное расписание. Есть выбор курса, группы, недели, даты, номера и времени пары.")
    add_label_paragraph(doc, "Безопасность публикации:", "черновая автогенерация остаётся у диспетчера, пока не нажата кнопка «Опубликовать».")

    doc.add_heading("4.2. Расписание", level=2)
    for text in (
        "Просмотр общей сетки по курсам, группам и неделям.",
        "Генерация с нуля, фиксация Конструктора или фиксация прошлой автогенерации.",
        "Отображение прогресса и возможность отмены длительного запуска.",
        "Экспорт текущего результата в Excel и PDF, отдельная публикация студентам.",
    ):
        add_list_item(doc, text, bullet_num)
    add_screenshot(doc, "01_schedule.png", "Рисунок 1. Рабочий экран диспетчерского расписания.")

    doc.add_heading("4.3. Конструктор", level=2)
    add_label_paragraph(doc, "Назначение:", "ручная правка сетки после генерации либо создание варианта с нуля.")
    add_label_paragraph(doc, "Контроль:", "при выборе занятия интерфейс предупреждает о занятости подгруппы, преподавателя и выполненной нагрузке.")
    add_label_paragraph(doc, "Связь с решателем:", "сохранённый Конструктор можно использовать как фиксированную основу следующей генерации.")
    add_screenshot(doc, "02_constructor.png", "Рисунок 2. Конструктор до копирования результата автогенерации.")

    doc.add_heading("4.4. Преподаватели", level=2)
    for text in (
        "Поиск по ФИО и ответственности за кабинет, фильтр по площадке.",
        "Карточка преподавателя: ФИО, закреплённый кабинет, приоритет площадок и рабочий график.",
        "Массовое выделение найденных и назначение рабочего времени, площадки или кабинета сразу нескольким преподавателям.",
        "Недоступность по диапазону дат учитывается генератором и матрицей занятости.",
    ):
        add_list_item(doc, text, bullet_num)
    add_screenshot(doc, "03_teachers.png", "Рисунок 3. Справочник преподавателей и массовое выделение.")

    doc.add_heading("4.5. Группы", level=2)
    for text in (
        "Название, количество подгрупп, численность и основной корпус.",
        "Рабочий период внутри семестра и допустимые дни/диапазоны пар.",
        "Кнопка «Выделить все» и массовое назначение рабочего времени выбранным группам.",
        "Численность используется при подборе аудитории по вместимости.",
    ):
        add_list_item(doc, text, bullet_num)
    add_screenshot(doc, "04_groups.png", "Рисунок 4. Справочник групп с выделением и параметрами.")

    doc.add_heading("4.6. Пары / занятия", level=2)
    for text in (
        "Дисциплина связана с группой, подгруппой, преподавателем и объёмом часов.",
        "Поддерживаются обычные занятия, ЛПЗ, УП-блоки и чётность недель.",
        "Можно ограничить площадки, вместимость, оборудование и тип аудитории.",
        "Закреплённый кабинет является приоритетом; при разрешённой автозамене конфликт не блокирует подбор другого совместимого кабинета.",
    ):
        add_list_item(doc, text, bullet_num)
    add_screenshot(doc, "08_lesson_room.png", "Рисунок 5. Требования к аудитории в карточке занятия.")

    doc.add_heading("4.7. Аудиторный фонд", level=2)
    add_label_paragraph(doc, "Карточка аудитории:", "номер/название, корпус, тип, вместимость, оборудование и признак доступности.")
    add_label_paragraph(doc, "Типы:", "по умолчанию 1 — лекционная, 2 — мастерская со станками, 3 — компьютерный класс; разрешены собственные коды и описания.")
    add_label_paragraph(doc, "Автозамена:", "сначала проверяются тип, корпус, вместимость и оборудование, затем предпочтения. При занятости закреплённого кабинета выбирается свободный совместимый.")
    add_label_paragraph(doc, "Отчёт:", "после генерации показывает число назначений, автозамен, занятий без кабинета и причину каждой замены.")
    add_screenshot(doc, "06_rooms_types.png", "Рисунок 6. Редактируемый справочник типов аудиторий.")

    doc.add_heading("4.8. Данные и контроль", level=2)
    data_rows = [
        ("Импорт Excel", "Предпросмотр «Вклеек», выбор семестра, предупреждения, аудит и применение с резервной версией."),
        ("Аудит", "Ошибки, предупреждения и информационные замечания до запуска решателя."),
        ("Учёт часов", "План, поставлено, зачтено, остаток, проценты, недели и конкретные даты пар."),
        ("Замены", "Больной/отсутствующий преподаватель, заменяющий, дата, пара, часы и причина."),
        ("Занятость", "Матрица преподаватель × день × пара с кабинетом и отметкой замены."),
        ("Недоступность", "Отпуск, больничный, методический день или другой диапазон недоступности преподавателя."),
        ("История", "До 50 последних версий, причина изменения и восстановление выбранной версии."),
    ]
    add_table(doc, ["Вкладка", "Функция"], data_rows, [2200, 7160], font_size=9.0)
    add_screenshot(doc, "05_hours.png", "Рисунок 7. Учёт часов по группам с недельной детализацией.")

    doc.add_heading("4.9. Настройки", level=2)
    for text in (
        "Даты начала и окончания семестра определяют календарную сетку и количество недель.",
        "Генерация запускается с нуля либо с фиксацией Конструктора/предыдущей автогенерации.",
        "Недоступные дни можно задать одной группе, всем группам, диапазоном или перечнем дат.",
        "Скрытое меню ROOT открывает детальные параметры CP-SAT и применяется при следующей регенерации.",
    ):
        add_list_item(doc, text, bullet_num)
    add_screenshot(doc, "07_root.png", "Рисунок 8. Начало меню параметров решателя ROOT.")
    doc.add_page_break()

    doc.add_heading("5. Как работает генератор", level=1)
    generator_steps = [
        "Данные нормализуются: проверяются идентификаторы, ссылки, даты и обязательные поля.",
        "Аудит выявляет недоступность, дефицит часов, ошибочные кабинеты и противоречащие требования.",
        "Семестр делится на недельные задачи. Это быстрее и устойчивее, чем одна огромная модель.",
        "CP-SAT ищет расписание без пересечений групп, подгрупп и преподавателей и соблюдает рабочее время.",
        "При сложной неделе возможен повтор с ослаблением необязательных условий.",
        "После первого допустимого решения отдельная фаза улучшает окна и другие мягкие показатели в заданный лимит.",
        "Готовая сетка передаётся распределителю аудиторий и далее в отчёты качества, часов и публикацию.",
    ]
    for item in generator_steps:
        add_list_item(doc, item, decimal_num)

    doc.add_heading("5.1. Жёсткие и мягкие условия", level=2)
    hard_soft_rows = [
        ("Жёсткие", "Пересечения, недоступность, рабочее время, верхние пределы, запрещённые окна при HARD. Нарушить нельзя — иначе решения нет."),
        ("Мягкие", "Окна, поздние пары, распределение предмета, число учебных дней. Нарушение возможно, но увеличивает штраф."),
        ("Веса", "Показывают solver-у, какое мягкое нарушение важнее. Больший вес означает более высокий приоритет."),
    ]
    add_table(doc, ["Тип", "Смысл"], hard_soft_rows, [1800, 7560], font_size=9.3)
    add_callout(
        doc,
        "Практическое правило",
        "Сначала получить допустимое расписание с мягкими ограничениями. Затем включать оптимизацию и усиливать веса. Не превращать все пожелания в HARD одновременно.",
        fill="FFF8E8",
        accent=AMBER,
    )

    doc.add_heading("5.2. Подбор и замена кабинета", level=2)
    room_steps = [
        "Берётся корпус дня группы и требуемый тип аудитории.",
        "Отбрасываются неактивные, слишком маленькие и неоснащённые кабинеты.",
        "Закреплённый кабинет занятия и кабинет преподавателя получают приоритет.",
        "Если запрошенный кабинет занят, назначается следующий совместимый свободный.",
        "Если вариантов нет, занятие остаётся без кабинета и попадает в конфликтный отчёт.",
    ]
    for item in room_steps:
        add_list_item(doc, item, decimal_num)

    doc.add_heading("6. Меню ROOT: назначение и правила", level=1)
    add_label_paragraph(doc, "Как открыть:", "перейти в «Настройки», щёлкнуть вне поля ввода и набрать на клавиатуре root!.")
    add_label_paragraph(doc, "Как применить:", "изменить значения, нажать «Сохранить» и запустить новую генерацию.")
    add_label_paragraph(doc, "Как отменить эксперименты:", "кнопка «Сбросить к дефолтам» возвращает безопасный базовый набор.")
    add_callout(
        doc,
        "Осторожно",
        "ROOT влияет на скорость, память и разрешимость модели. Перед серией экспериментов рекомендуется сохранить версию данных и менять одну группу параметров за раз.",
        fill="FEF3F2",
        accent=RED,
    )

    category_titles = {
        "solver": "6.1. Solver: скорость и ресурсы",
        "hard_soft": "6.2. Жёсткие и мягкие переключатели",
        "weights": "6.3. Веса штрафов",
        "shape": "6.4. Размерности расписания",
    }
    category_notes = {
        "solver": "Эти параметры определяют время поиска, число потоков, память и глубину улучшения результата.",
        "hard_soft": "Переключатели меняют характер требований: запрет, мягкое предпочтение или полное отключение оптимизации.",
        "weights": "Вес 0 практически отключает соответствующее предпочтение. Увеличивать веса следует постепенно.",
        "shape": "Размерности задают допустимую форму учебного дня и распределение дисциплин по семестру.",
    }
    schema = solver.get("schema", [])
    values = solver.get("values", {})
    for category in ("solver", "hard_soft", "weights", "shape"):
        doc.add_page_break()
        doc.add_heading(category_titles[category], level=1)
        p = doc.add_paragraph(category_notes[category])
        p.paragraph_format.space_after = Pt(8)
        rows = []
        for field in [x for x in schema if x.get("category") == category]:
            key = field["key"]
            purpose = field.get("description", "")
            advice = ROOT_ADVICE.get(key, "Менять только после сравнения результатов.")
            rows.append((f"{field['label']}\n{key}", value_text(values.get(key, "—")), f"{purpose}\nСовет: {advice}"))
        add_table(doc, ["Параметр", "Сейчас", "Что даёт и как использовать"], rows, [2600, 1000, 5760], font_size=8.25)

    doc.add_page_break()
    doc.add_heading("7. Учёт часов и замены: логика", level=1)
    accounting_rows = [
        ("План", "Объём из действующих «Вклеек» по дисциплине, группе и преподавателю."),
        ("Поставлено", "Часы занятий, реально размещённых в расписании по конкретным датам."),
        ("Зачтено", "Фактические часы преподавателя после учёта замен и корректировок."),
        ("Осталось", "План минус зачтено; отрицательное значение означает перегруз."),
        ("Замена", "У отсутствующего часы вычитаются из зачёта, заменяющему добавляются; исходный план не переписывается."),
    ]
    add_table(doc, ["Показатель", "Правило расчёта"], accounting_rows, [2100, 7260], font_size=9.3)
    add_label_paragraph(doc, "Детализация:", "по нажатию на часы открываются даты, недели, номера пар, группы, дисциплины, преподаватели и кабинеты.")
    add_label_paragraph(doc, "Экспорт:", "универсальный Excel содержит сводные листы, а «Экспорт по образцу» строит книгу по структуре старого учёта часов.")

    doc.add_heading("8. Что показать на презентации", level=1)
    demo_steps = [
        "Открыть расписание и показать выбор курса/недели, Excel, PDF и публикацию.",
        "Открыть преподавателей: выделить несколько карточек и показать массовое рабочее время.",
        "Открыть занятие: показать чётность недели, тип аудитории и автозамену кабинета.",
        "Открыть аудиторный фонд: продемонстрировать типы 1/2/3 и возможность добавить свой.",
        "Открыть учёт часов: план, поставлено, зачтено, остаток, даты и экспорт по образцу.",
        "Показать журнал замен и объяснить перенос фактических часов.",
        "В Настройках набрать root! и объяснить разделы без изменения рабочих значений.",
    ]
    for item in demo_steps:
        add_list_item(doc, item, decimal_num)

    doc.add_heading("9. Ограничения текущей версии", level=1)
    limitations = [
        "Качество результата зависит от полноты входных данных: численности, кабинетов, рабочего времени и недоступности.",
        "Тип аудитории сейчас сопоставляется точно: тип 3 не считается автоматической заменой типа 1 без отдельного правила совместимости.",
        "После ручной правки, импорта или ROOT-изменений требуется явная регенерация.",
        "Перед публикацией диспетчер должен проверить аудит, часы, кабинеты и визуальную сетку.",
        "ROOT предназначен для ответственного пользователя; неверные HARD-настройки могут привести к отсутствию решения.",
    ]
    for item in limitations:
        add_list_item(doc, item, bullet_num)

    doc.add_heading("10. Следующие логичные улучшения", level=1)
    next_steps = [
        "Матрица совместимости типов: например, компьютерный класс может принимать лекцию, но мастерская — только профильные занятия.",
        "Приоритеты и запреты для отдельных кабинетов/дисциплин вместо только точного типа.",
        "Сценарное сравнение двух генераций по окнам, поздним парам, сменам корпусов и заполнению кабинетов.",
        "Ролевой доступ: диспетчер, администрация, преподаватель и студент.",
        "Автоматическая публикация выбранной версии по расписанию после подтверждения диспетчером.",
    ]
    for item in next_steps:
        add_list_item(doc, item, bullet_num)

    add_callout(
        doc,
        "Итог",
        "Система уже закрывает полный демонстрационный цикл: входные данные → аудит → генерация → ручная проверка → кабинеты → учёт часов → публикация и отчёты.",
        fill="EEF2FF",
        accent=BLUE,
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
